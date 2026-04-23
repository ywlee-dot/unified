"""문서 로딩 모듈"""

import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Generator, Literal

import chardet
from loguru import logger

DocType = Literal["institutions", "data_catalog", "analysis_cases"]


@dataclass
class Document:
    """문서 데이터 클래스"""
    content: str
    metadata: dict = field(default_factory=dict)
    source: str = ""
    doc_type: DocType | None = None

    def __post_init__(self):
        if self.source and not self.metadata.get("source"):
            self.metadata["source"] = self.source
        if self.doc_type and not self.metadata.get("doc_type"):
            self.metadata["doc_type"] = self.doc_type


class BaseLoader(ABC):
    """문서 로더 기본 클래스"""
    supported_extensions: list[str] = []

    @classmethod
    def can_load(cls, file_path: Path) -> bool:
        return file_path.suffix.lower() in cls.supported_extensions

    @abstractmethod
    def load(self, file_path: Path) -> Document | None:
        pass


class PDFLoader(BaseLoader):
    """PDF 문서 로더"""
    supported_extensions = [".pdf"]

    def load(self, file_path: Path) -> Document | None:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            content = "\n\n".join(text_parts)
            if not content.strip():
                logger.warning(f"PDF 파일에서 텍스트를 추출할 수 없음: {file_path}")
                return None

            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "page_count": len(reader.pages),
                }
            )
        except Exception as e:
            logger.error(f"PDF 로딩 실패 ({file_path}): {e}")
            return None


class HWPLoader(BaseLoader):
    """HWP 문서 로더 (olefile 사용)"""
    supported_extensions = [".hwp"]

    def load(self, file_path: Path) -> Document | None:
        try:
            import olefile

            ole = olefile.OleFileIO(str(file_path))

            # PrvText 스트림에서 텍스트 추출
            if ole.exists("PrvText"):
                encoded_text = ole.openstream("PrvText").read()
                content = encoded_text.decode("utf-16-le", errors="ignore")
                content = content.replace("\x00", "")
            else:
                # BodyText 섹션들에서 추출 시도
                text_parts = []
                for stream in ole.listdir():
                    stream_path = "/".join(stream)
                    if "BodyText" in stream_path or "Section" in stream_path:
                        try:
                            data = ole.openstream(stream).read()
                            # 텍스트 추출 시도
                            detected = chardet.detect(data)
                            if detected["encoding"]:
                                text = data.decode(detected["encoding"], errors="ignore")
                                text_parts.append(text)
                        except Exception:
                            pass
                content = "\n".join(text_parts)

            ole.close()

            if not content.strip():
                logger.warning(f"HWP 파일에서 텍스트를 추출할 수 없음: {file_path}")
                return None

            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": "hwp",
                }
            )
        except Exception as e:
            logger.error(f"HWP 로딩 실패 ({file_path}): {e}")
            return None


class DOCXLoader(BaseLoader):
    """DOCX 문서 로더"""
    supported_extensions = [".docx"]

    def load(self, file_path: Path) -> Document | None:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 테이블 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            content = "\n\n".join(text_parts)
            if not content.strip():
                logger.warning(f"DOCX 파일에서 텍스트를 추출할 수 없음: {file_path}")
                return None

            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": "docx",
                }
            )
        except Exception as e:
            logger.error(f"DOCX 로딩 실패 ({file_path}): {e}")
            return None


class ExcelLoader(BaseLoader):
    """Excel 문서 로더"""
    supported_extensions = [".xlsx", ".xls", ".csv"]

    def load(self, file_path: Path) -> Document | None:
        try:
            import pandas as pd

            if file_path.suffix.lower() == ".csv":
                # CSV 인코딩 자동 감지
                with open(file_path, "rb") as f:
                    raw_data = f.read()
                    detected = chardet.detect(raw_data)
                    encoding = detected.get("encoding", "utf-8")
                df = pd.read_csv(file_path, encoding=encoding)
            else:
                # Excel 파일의 모든 시트 읽기
                excel_file = pd.ExcelFile(file_path)
                dfs = []
                for sheet_name in excel_file.sheet_names:
                    sheet_df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_df["_sheet_name"] = sheet_name
                    dfs.append(sheet_df)
                df = pd.concat(dfs, ignore_index=True) if dfs else pd.DataFrame()

            if df.empty:
                logger.warning(f"Excel 파일이 비어있음: {file_path}")
                return None

            # DataFrame을 텍스트로 변환
            text_parts = []

            # 컬럼 정보
            columns = [str(c) for c in df.columns if not str(c).startswith("_")]
            text_parts.append(f"컬럼: {', '.join(columns)}")
            text_parts.append("")

            # 각 행을 텍스트로 변환
            for idx, row in df.iterrows():
                row_text_parts = []
                for col in df.columns:
                    if str(col).startswith("_"):
                        continue
                    val = row[col]
                    if pd.notna(val):
                        row_text_parts.append(f"{col}: {val}")
                if row_text_parts:
                    text_parts.append(" | ".join(row_text_parts))

            content = "\n".join(text_parts)

            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": file_path.suffix.lower().replace(".", ""),
                    "row_count": len(df),
                    "column_count": len(columns),
                }
            )
        except Exception as e:
            logger.error(f"Excel 로딩 실패 ({file_path}): {e}")
            return None


class TXTLoader(BaseLoader):
    """텍스트 문서 로더"""
    supported_extensions = [".txt", ".md"]

    def load(self, file_path: Path) -> Document | None:
        try:
            # 인코딩 자동 감지
            with open(file_path, "rb") as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get("encoding", "utf-8")

            content = raw_data.decode(encoding, errors="ignore")

            if not content.strip():
                logger.warning(f"텍스트 파일이 비어있음: {file_path}")
                return None

            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": file_path.suffix.lower().replace(".", ""),
                }
            )
        except Exception as e:
            logger.error(f"텍스트 로딩 실패 ({file_path}): {e}")
            return None


class DocumentLoaderFactory:
    """문서 로더 팩토리"""

    loaders: list[type[BaseLoader]] = [
        PDFLoader,
        HWPLoader,
        DOCXLoader,
        ExcelLoader,
        TXTLoader,
    ]

    @classmethod
    def get_loader(cls, file_path: Path) -> BaseLoader | None:
        for loader_class in cls.loaders:
            if loader_class.can_load(file_path):
                return loader_class()
        return None

    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        extensions = []
        for loader_class in cls.loaders:
            extensions.extend(loader_class.supported_extensions)
        return extensions

    @classmethod
    def load_file(cls, file_path: Path, doc_type: DocType | None = None) -> Document | None:
        loader = cls.get_loader(file_path)
        if loader is None:
            logger.warning(f"지원하지 않는 파일 형식: {file_path}")
            return None

        doc = loader.load(file_path)
        if doc and doc_type:
            doc.doc_type = doc_type
            doc.metadata["doc_type"] = doc_type
        return doc

    @classmethod
    def load_directory(
        cls,
        dir_path: Path,
        doc_type: DocType | None = None,
        recursive: bool = True
    ) -> Generator[Document, None, None]:
        """디렉토리의 모든 문서 로딩"""
        if not dir_path.exists():
            logger.warning(f"디렉토리가 존재하지 않음: {dir_path}")
            return

        supported_ext = cls.get_supported_extensions()

        if recursive:
            files = list(dir_path.rglob("*"))
        else:
            files = list(dir_path.glob("*"))

        for file_path in files:
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in supported_ext:
                continue

            # 기관명 추출 (institutions 폴더의 경우)
            institution_name = None
            if doc_type == "institutions":
                # institutions/기관명/파일.ext 구조
                rel_path = file_path.relative_to(dir_path)
                if len(rel_path.parts) > 1:
                    institution_name = rel_path.parts[0]

            doc = cls.load_file(file_path, doc_type)
            if doc:
                if institution_name:
                    doc.metadata["institution_name"] = institution_name
                yield doc
